import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog
import fitz  # PyMuPDF, for PDF processing
import re # For table formatting
import json # For parsing JSON responses
# Try to import specific fitz error, fall back if not found (e.g. older PyMuPDF)
try:
    from fitz.errors import FitzError
except ImportError:
    FitzError = Exception # Fallback to generic Exception if specific error not found
import os
import shutil
import google.generativeai as genai
from google.api_core import exceptions as google_exceptions
# Import specific GenAI exceptions if needed, e.g., genai.types.BlockedPromptException
try:
    from google.generativeai.types import BlockedPromptException, StopCandidateException
except ImportError:
    BlockedPromptException = Exception # Fallback
    StopCandidateException = Exception # Fallback

from PIL import Image, ImageTk, UnidentifiedImageError # Pillow for image handling
import docx # For downloading history
from docx.shared import RGBColor # For coloring text in Word
from dotenv import load_dotenv # For loading .env files
import traceback # For detailed error logging

class Tooltip:
    """
    Create a tooltip for a given widget.
    """
    def __init__(self, widget, text_callback):
        self.widget = widget
        self.text_callback = text_callback
        self.tip_window = None
        self.widget.bind("<Enter>", self.show_tip)
        self.widget.bind("<Leave>", self.hide_tip)

    def show_tip(self, event=None):
        """Display tooltip"""
        if self.tip_window or not hasattr(self, 'text_callback'):
            return

        text = self.text_callback()
        if not text:
            return

        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 1

        self.tip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")

        label = tk.Label(tw, text=text, justify=tk.LEFT,
                         background="#ffffe0", relief=tk.SOLID, borderwidth=1,
                         font=("tahoma", "8", "normal"))
        label.pack(ipadx=1)

    def hide_tip(self, event=None):
        """Hide tooltip"""
        if self.tip_window:
            self.tip_window.destroy()
        self.tip_window = None

class ComponentComparatorAI:
    """
    Main application class for the Component Comparator AI.
    Manages the UI, file loading, PDF processing, AI interaction,
    and conversation history.
    """
    def __init__(self, root):
        """
        Initializes the application UI and internal state.
        """
        self.root = root
        self.root.title("Component Comparator AI")
        self.root.geometry("850x870")

        load_dotenv()

        self.spec_sheet_1_path = None; self.spec_sheet_1_text = None; self.spec_sheet_1_image_paths = []
        self.mfg_pn_var_1 = tk.StringVar()
        self.spec_sheet_2_path = None; self.spec_sheet_2_text = None; self.spec_sheet_2_image_paths = []
        self.mfg_pn_var_2 = tk.StringVar()

        self.model = None; self.chat_session = None; self.conversation_log = []; self.ai_history = []
        self.api_key_configured = False; self.model_options_list = []
        self.placeholder_text = "Select AI Model (after loading files)"; self.model_initializing = False
        self.start_comparison_button = None
        self.upload_image_button = None
        self.pending_user_image_path = None
        self.pending_user_image_pil = None
        self.translate_to_chinese_var = tk.BooleanVar(value=False)


        self.temp_image_dir = "temp_images"; self._create_temp_image_dir()
        self._setup_ui(root); self._configure_ai()

        self.update_conversation_history(
            "System: Welcome! Please load two PDF specification sheets to compare. " +
            "After loading both files, you will be able to select an AI model.", role="system"
        )
        self._update_ui_for_ai_status(api_key_configured=self.api_key_configured, model_initialized=False)


    def _create_temp_image_dir(self):
        if not os.path.exists(self.temp_image_dir):
            try: os.makedirs(self.temp_image_dir)
            except OSError as e: print(f"Critical Error creating temp dir {self.temp_image_dir}: {e}"); self.update_conversation_history(f"System: Error creating temp folder: {e}", role="error")

    def _setup_ui(self, root):
        current_row = 0
        # File 1 & MFG P/N 1
        self.spec_sheet_1_label = ttk.Label(root, text="File 1: None")
        self.spec_sheet_1_label.grid(row=current_row, column=0, padx=10, pady=2, sticky="w")
        self.load_spec_sheet_1_button = ttk.Button(root, text="Load Spec Sheet 1", command=self.load_spec_sheet_1)
        self.load_spec_sheet_1_button.grid(row=current_row, column=1, padx=5, pady=2, sticky="ew")
        current_row += 1
        self.mfg_pn_label_1 = ttk.Label(root, text="MFG P/N 1:")
        self.mfg_pn_label_1.grid(row=current_row, column=0, sticky=tk.W, padx=10, pady=2)
        self.mfg_pn_entry_1 = ttk.Entry(root, textvariable=self.mfg_pn_var_1, width=30)
        self.mfg_pn_entry_1.grid(row=current_row, column=1, sticky=tk.EW, padx=5, pady=2)
        self.mfg_pn_entry_1.bind("<FocusOut>", self._handle_mfg_pn1_entry_change)
        self.mfg_pn_entry_1.bind("<Return>", self._handle_mfg_pn1_entry_change)
        current_row += 1
        # File 2 & MFG P/N 2
        self.spec_sheet_2_label = ttk.Label(root, text="File 2: None")
        self.spec_sheet_2_label.grid(row=current_row, column=0, padx=10, pady=2, sticky="w")
        self.load_spec_sheet_2_button = ttk.Button(root, text="Load Spec Sheet 2", command=self.load_spec_sheet_2)
        self.load_spec_sheet_2_button.grid(row=current_row, column=1, padx=5, pady=2, sticky="ew")
        current_row += 1
        self.mfg_pn_label_2 = ttk.Label(root, text="MFG P/N 2:")
        self.mfg_pn_label_2.grid(row=current_row, column=0, sticky=tk.W, padx=10, pady=2)
        self.mfg_pn_entry_2 = ttk.Entry(root, textvariable=self.mfg_pn_var_2, width=30)
        self.mfg_pn_entry_2.grid(row=current_row, column=1, sticky=tk.EW, padx=5, pady=2)
        self.mfg_pn_entry_2.bind("<FocusOut>", self._handle_mfg_pn2_entry_change)
        self.mfg_pn_entry_2.bind("<Return>", self._handle_mfg_pn2_entry_change)
        current_row += 1

        # Model Selection
        ttk.Label(root, text="Select AI Model:").grid(row=current_row, column=0, padx=10, pady=5, sticky="w")
        self.model_var = tk.StringVar()
        self.model_combobox = ttk.Combobox(root, textvariable=self.model_var, width=50)
        self.model_options_list = ["models/gemini-1.0-pro-vision-latest", "models/gemini-pro-vision","models/gemini-1.5-flash-latest", "models/gemini-1.5-flash","models/gemini-1.5-flash-002", "models/gemini-1.5-flash-8b","models/gemini-1.5-flash-8b-001", "models/gemini-1.5-flash-8b-latest","models/gemini-2.5-flash-preview-04-17", "models/gemini-2.5-flash-preview-05-20","models/gemini-2.5-flash-preview-04-17-thinking", "models/gemini-2.0-flash-exp","models/gemini-2.0-flash", "models/gemini-2.0-flash-001","models/gemini-2.0-flash-exp-image-generation", "models/gemini-2.0-flash-lite-001","models/gemini-2.0-flash-lite", "models/gemini-2.0-flash-lite-preview-02-05","models/gemini-2.0-flash-lite-preview", "models/gemini-2.0-flash-thinking-exp-01-21","models/gemini-2.0-flash-thinking-exp", "models/gemini-2.0-flash-thinking-exp-1219","models/learnlm-2.0-flash-experimental", "models/gemma-3-1b-it","models/gemma-3-4b-it", "models/gemma-3-12b-it","models/gemma-3-27b-it", "models/gemma-3n-e4b-it"]
        self.model_combobox['values'] = self.model_options_list
        self.model_combobox.set(self.placeholder_text); self.model_combobox.state(["disabled"])
        self.model_combobox.grid(row=current_row, column=1, padx=5, pady=5, sticky="ew")
        self.model_combobox.bind("<<ComboboxSelected>>", self._on_model_selected)
        def get_tooltip_text(): return self.model_combobox.get() if self.model_combobox.get() != self.placeholder_text else None
        self.combobox_tooltip = Tooltip(self.model_combobox, get_tooltip_text)
        current_row += 1
        # Conversation History
        ttk.Label(root, text="Conversation History:").grid(row=current_row, column=0, columnspan=2, padx=10, pady=(10,0), sticky="w")
        current_row += 1
        self.conversation_history = scrolledtext.ScrolledText(root, wrap=tk.WORD, height=10, width=80)
        self.conversation_history.grid(row=current_row, column=0, columnspan=2, padx=10, pady=5, sticky="nsew")
        self.conversation_history.tag_configure("user_message", foreground="blue", font=('Arial', 10))
        self.conversation_history.tag_configure("ai_message", foreground="#008800", font=('Arial', 10))
        self.conversation_history.tag_configure("system_message", foreground="#550055", font=('Arial', 10, 'italic'))
        self.conversation_history.tag_configure("error_message", foreground="red", font=('Arial', 10, 'bold'))
        self.conversation_history.config(state=tk.DISABLED)
        root.grid_rowconfigure(current_row, weight=1)
        current_row += 1
        # Comparison Treeview
        self.treeview_frame = ttk.LabelFrame(root, text="Comparison Details")
        self.treeview_frame.grid(row=current_row, column=0, columnspan=2, sticky="nsew", padx=10, pady=5)
        columns = ("parameter", "component1", "component2", "notes")
        self.comparison_treeview = ttk.Treeview(self.treeview_frame, columns=columns, show="headings", height=8)
        self.comparison_treeview.heading("parameter", text="Parameter"); self.comparison_treeview.column("parameter", anchor=tk.W, width=150)
        self.comparison_treeview.heading("component1", text="Comp 1 Val"); self.comparison_treeview.column("component1", anchor=tk.W, width=200)
        self.comparison_treeview.heading("component2", text="Comp 2 Val"); self.comparison_treeview.column("component2", anchor=tk.W, width=200)
        self.comparison_treeview.heading("notes", text="Notes"); self.comparison_treeview.column("notes", anchor=tk.W, width=250)
        vsb = ttk.Scrollbar(self.treeview_frame, orient="vertical", command=self.comparison_treeview.yview)
        hsb = ttk.Scrollbar(self.treeview_frame, orient="horizontal", command=self.comparison_treeview.xview)
        self.comparison_treeview.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        self.comparison_treeview.grid(row=0, column=0, sticky="nsew"); vsb.grid(row=0, column=1, sticky="ns"); hsb.grid(row=1, column=0, sticky="ew")
        self.treeview_frame.grid_rowconfigure(0, weight=1); self.treeview_frame.grid_columnconfigure(0, weight=1)
        root.grid_rowconfigure(current_row, weight=1)
        current_row += 1

        # --- User Input Controls Frame ---
        self.user_input_controls_frame = ttk.Frame(root)
        self.user_input_controls_frame.grid(row=current_row, column=0, columnspan=2, sticky="ew", padx=5, pady=0)
        self.user_input_controls_frame.columnconfigure(0, weight=1)
        self.user_input_controls_frame.columnconfigure(1, weight=0)
        self.user_input_controls_frame.columnconfigure(2, weight=0)
        self.user_input_controls_frame.columnconfigure(3, weight=0)

        ttk.Label(self.user_input_controls_frame, text="Your Message:").grid(row=0, column=0, columnspan=4, padx=5, pady=(5,0), sticky="w")

        self.user_input_entry = ttk.Entry(self.user_input_controls_frame, width=70)
        self.user_input_entry.grid(row=1, column=0, padx=(5,0), pady=5, sticky="ew")

        self.translate_to_chinese_checkbutton = ttk.Checkbutton(
            self.user_input_controls_frame,
            text="AI replies in Chinese",
            #variable=self.translate_to_chinese_var,  # 綁定變數
            command=self._handle_translate_chinese_checkbox_change, # Removed variable, added command
            onvalue=True,
            offvalue=False
        )
        self.translate_to_chinese_checkbutton.grid(row=1, column=1, padx=5, pady=5, sticky="e")
        # Set initial visual state of the checkbutton based on the BooleanVar's current value
        if self.translate_to_chinese_var.get():
            self.translate_to_chinese_checkbutton.state(['selected'])
        else:
            self.translate_to_chinese_checkbutton.state(['!selected'])

        self.upload_image_button = ttk.Button(self.user_input_controls_frame, text="Attach Image", command=self.on_upload_image)
        self.upload_image_button.grid(row=1, column=2, padx=5, pady=5, sticky="e")

        self.send_button = ttk.Button(self.user_input_controls_frame, text="Send", command=self.send_user_query)
        self.send_button.grid(row=1, column=3, padx=(0,5), pady=5, sticky="e")
        current_row += 1

        # --- Action Buttons Frame ---
        self.action_buttons_frame = ttk.Frame(root)
        self.action_buttons_frame.grid(row=current_row, column=0, columnspan=2, sticky="ew", padx=5, pady=5)
        self.action_buttons_frame.columnconfigure(0, weight=1)
        self.action_buttons_frame.columnconfigure(1, weight=1)
        self.action_buttons_frame.columnconfigure(2, weight=1)
        self.start_comparison_button = ttk.Button(self.action_buttons_frame, text="Start Detailed Comparison", command=self.on_start_detailed_comparison, state=tk.DISABLED)
        self.start_comparison_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5, pady=5)
        self.download_history_button = ttk.Button(self.action_buttons_frame, text="Download History", command=self.download_history)
        self.download_history_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5, pady=5)
        self.clear_all_button = ttk.Button(self.action_buttons_frame, text="Clear All", command=self.clear_all)
        self.clear_all_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5, pady=5)

        root.grid_columnconfigure(0, weight=1) # Label column, less expansion needed
        root.grid_columnconfigure(1, weight=1) # Input widgets column, allow expansion

    def _handle_translate_chinese_checkbox_change(self):
        # Checkbutton's state() method returns a tuple of state flags.
        # 'selected' is present if checked.
        is_selected = 'selected' in self.translate_to_chinese_checkbutton.state()
        self.translate_to_chinese_var.set(is_selected)
        # The existing trace on self.translate_to_chinese_var should fire after .set()

    def _handle_mfg_pn1_entry_change(self, event=None):
        current_text = self.mfg_pn_entry_1.get()
        self.mfg_pn_var_1.set(current_text)

    def _handle_mfg_pn2_entry_change(self, event=None):
        current_text = self.mfg_pn_entry_2.get()
        self.mfg_pn_var_2.set(current_text)

    def _parse_initial_analysis_response(self, response_text: str) -> dict:
        """Parses the structured AI response from the initial analysis."""
        data = {
            "component1_type": "Unknown",
            "component2_type": "Unknown",
            "functionally_similar": "Unknown",
            "is_similar_flag": False,  # Default to False
            "mfg_pn1": "Not Found",
            "mfg_pn2": "Not Found"
        }
        try:
            # Attempt to parse as JSON first
            # Remove potential markdown backticks from JSON string
            cleaned_response_text = response_text.strip()
            if cleaned_response_text.startswith("```json"):
                cleaned_response_text = cleaned_response_text[7:]
            if cleaned_response_text.startswith("```"):
                cleaned_response_text = cleaned_response_text[3:]
            if cleaned_response_text.endswith("```"):
                cleaned_response_text = cleaned_response_text[:-3]

            parsed_json = json.loads(cleaned_response_text.strip())
            data["component1_type"] = parsed_json.get("Component1_Type", "Unknown")
            data["component2_type"] = parsed_json.get("Component2_Type", "Unknown")
            similarity_text = parsed_json.get("Functionally_Similar", "Unknown")
            data["functionally_similar"] = similarity_text
            if isinstance(similarity_text, str) and (similarity_text.lower().startswith("yes") or similarity_text.startswith("是")):
                data["is_similar_flag"] = True
            data["mfg_pn1"] = parsed_json.get("MFG_PN1", "Not Found")
            data["mfg_pn2"] = parsed_json.get("MFG_PN2", "Not Found")

            # Ensure "Not Found" from JSON still results in empty string for P/N variables if needed by downstream logic
            # (Current downstream logic in send_to_ai seems to handle "Not Found" correctly by setting P/N var to "" or the value)

        except json.JSONDecodeError:
            # Fallback to original line-by-line parsing if JSON parsing fails
            lines = response_text.split('\n')
            for line in lines:
                if line.startswith("Component1_Type:"):
                    data["component1_type"] = line.split(":", 1)[1].strip()
                elif line.startswith("Component2_Type:"):
                    data["component2_type"] = line.split(":", 1)[1].strip()
                elif line.startswith("Functionally_Similar:"):
                    similarity_text = line.split(":", 1)[1].strip()
                    data["functionally_similar"] = similarity_text
                    if similarity_text.lower().startswith("yes") or similarity_text.startswith("是"):
                        data["is_similar_flag"] = True
                elif line.startswith("MFG_PN1:"):
                    data["mfg_pn1"] = line.split(":", 1)[1].strip()
                elif line.startswith("MFG_PN2:"):
                    data["mfg_pn2"] = line.split(":", 1)[1].strip()
        return data

    def on_upload_image(self):
        filetypes = [('Image files', '*.png *.jpg *.jpeg *.bmp *.gif *.webp'), ('All files', '*.*')]
        filepath = filedialog.askopenfilename(title="Select an Image for AI Analysis", filetypes=filetypes)
        if filepath:
            try:
                img = Image.open(filepath)
                self.pending_user_image_path = filepath
                self.pending_user_image_pil = img
                image_name = os.path.basename(filepath)
                self.update_conversation_history(f"System: Image '{image_name}' attached. It will be sent with your next message.", role="system")
            except FileNotFoundError:
                self.update_conversation_history(f"System: Error - Image file not found at {filepath}", role="error")
                self.pending_user_image_path = None; self.pending_user_image_pil = None
            except UnidentifiedImageError:
                self.update_conversation_history(f"System: Error - Cannot identify image file. Not a valid image format? File: {filepath}", role="error")
                self.pending_user_image_path = None; self.pending_user_image_pil = None
            except Exception as e:
                self.update_conversation_history(f"System: Error processing image {filepath}: {e}", role="error")
                self.pending_user_image_path = None; self.pending_user_image_pil = None

    def on_start_detailed_comparison(self):
        self.update_conversation_history("System: 'Start Detailed Comparison' initiated...", role="system")
        if hasattr(self, 'start_comparison_button'):
            self.start_comparison_button.config(state=tk.DISABLED)

        try:
            if not self.model:
                self.update_conversation_history("System: No AI model initialized. Please select a model.", role="error")
                return
            if not (self.spec_sheet_1_text and self.spec_sheet_2_text):
                self.update_conversation_history("System: Both spec sheets must be loaded and processed.", role="error")
                return

            if hasattr(self, 'root'): self.root.update_idletasks()
            mfg_pn1 = self.mfg_pn_var_1.get() if hasattr(self, 'mfg_pn_var_1') and self.mfg_pn_var_1.get() else "N/A"
            mfg_pn2 = self.mfg_pn_var_2.get() if hasattr(self, 'mfg_pn_var_2') and self.mfg_pn_var_2.get() else "N/A"

            # Stage 1: Fetch Parameters
            self.update_conversation_history(f"System: Stage 1: Fetching relevant parameters for {mfg_pn1} vs {mfg_pn2}...", role="system")

            stage1_prompt_text = (
                f"Based on the two electronic components identified by MFG P/N 1: {mfg_pn1} and MFG P/N 2: {mfg_pn2}, "
                "please list the crucial electrical and physical parameters relevant for comparing these specific component types. "
                "Focus on parameters typically found in datasheets that are essential for electrical engineers to make a selection. "
                "Only list the parameter names, separated by commas."
            )
            stage1_user_prompt_for_history = f"User: Request for key parameters for detailed comparison of {mfg_pn1} vs {mfg_pn2}."

            parameters_response_text = self.send_to_ai(
                [stage1_prompt_text],
                is_initial_analysis=False, # This is part of detailed comparison, not initial one
                user_prompt_for_history=stage1_user_prompt_for_history
            )

            if not parameters_response_text or parameters_response_text.startswith("AI Error:") or "empty/no content" in parameters_response_text :
                self.update_conversation_history("System: Stage 1 Failed: Could not fetch relevant parameters from AI. Aborting detailed comparison.", role="error")
                if not parameters_response_text: # send_to_ai might return None
                     self.update_conversation_history("System: Failed to get parameters from AI (no response).", role="error")
                return

            # Clean up parameter list - remove potential numbering, newlines, and make it a comma separated string
            identified_parameters = re.sub(r"^\s*[\d.\-\s)]+\s*", "", parameters_response_text.strip(), flags=re.MULTILINE) # Remove leading numbers/bullets
            identified_parameters = identified_parameters.replace('\n', ', ').replace(';','_').replace('，',',').strip() # Replace newlines/other separators with commas
            identified_parameters = ", ".join(filter(None, [p.strip() for p in identified_parameters.split(',')])) # Ensure clean comma separation

            if not identified_parameters:
                self.update_conversation_history("System: Stage 1 Warning: AI did not return any parameters. Proceeding with general comparison.", role="system")
                # Fallback: if AI returns no params, provide a generic list or let stage 2 proceed without specific guidance.
                # For now, we'll let stage 2 proceed and it will ask for general differences.
            else:
                self.update_conversation_history(f"System: AI identified parameters: {identified_parameters}", role="system")

            # Stage 2: Fetch Detailed Differences Based on Parameters
            self.update_conversation_history(f"System: Stage 2: Fetching detailed differences for {mfg_pn1} vs {mfg_pn2} based on identified parameters...", role="system")

            datasheet_sourcing_instruction = (
                "IMPORTANT INSTRUCTIONS FOR AI RESPONSE:\n"
                "- Base your answers primarily on the information extracted from the provided component datasheets "
                "(text, images, and context from previous turns).\n"
                "- If the datasheets lack specific information to answer a point, you may use your general knowledge.\n"
                "- If you use general knowledge, you MUST explicitly state for which points the information was not found "
                "in the datasheets and that your answer for those points is based on general understanding.\n"
            )

            stage2_prompt_parts = [
                f"You are comparing two electronic components: MFG P/N 1: {mfg_pn1} and MFG P/N 2: {mfg_pn2}.\n\n"
                "--- COMPONENT 1 DATASHEET TEXT START ---\n"
                f"{self.spec_sheet_1_text}\n"
                "--- COMPONENT 1 DATASHEET TEXT END ---\n\n"
                "--- COMPONENT 2 DATASHEET TEXT START ---\n"
                f"{self.spec_sheet_2_text}\n"
                "--- COMPONENT 2 DATASHEET TEXT END ---\n\n",
                datasheet_sourcing_instruction,
            ]

            if identified_parameters:
                 stage2_prompt_parts.append(f"Focus on the following parameters that were identified for these components: {identified_parameters}.\n")
            else:
                 stage2_prompt_parts.append("Focus on crucial electrical and physical parameters relevant for comparing these specific component types.\n")

            stage2_prompt_parts.extend([
                "Based PRIMARILY on the provided datasheet texts above, please perform the following:",
                "1. List all key specification differences (especially considering the parameters above if provided) in a clear, concise markdown table format. Ensure the table includes columns for Parameter, Value for Component 1, and Value for Component 2. Include a 'Notes' or 'Difference' column if applicable.",
                "2. Explicitly state their full Operating Temperature ranges (e.g., -40°C to 125°C).",
                "3. Assess SMT Compatibility: Can Component 2's package (based on its description in provided text/images - though prioritize the full texts now re-provided) likely be SMT'd onto Component 1's typical PCB footprint? Consider common package names and pin counts. State any assumptions clearly.",
                "4. Package size including leads for two parts"
            ])

            stage2_user_prompt_for_history = (
                f"User: Request for detailed specification differences, temp ranges, and SMT compatibility "
                f"for {mfg_pn1} vs {mfg_pn2}"
                f"{', based on parameters: ' + identified_parameters if identified_parameters else '.'}"
            )

            detailed_comparison_response_text = self.send_to_ai(
                stage2_prompt_parts,
                is_initial_analysis=False,
                user_prompt_for_history=stage2_user_prompt_for_history
            )
            print("DEBUG:",detailed_comparison_response_text)
            if detailed_comparison_response_text and not detailed_comparison_response_text.startswith("AI Error:"):
                self._populate_comparison_treeview(detailed_comparison_response_text)
            else:
                if not detailed_comparison_response_text:
                    self.update_conversation_history("System: Stage 2 Failed: Failed to get detailed comparison from AI (no response).", role="error")
                # Error message already logged by send_to_ai if it starts with "AI Error:"
        finally:
            if self.model and hasattr(self, 'start_comparison_button'):
                self.start_comparison_button.config(state=tk.NORMAL)


    def _parse_markdown_table(self, markdown_text: str) -> tuple[dict or None, int]:
        # Filter out empty lines and strip whitespace
        # Keep track of original lines to count consumption accurately based on input structure
        original_lines = markdown_text.splitlines()
        processed_lines_info = [] # Stores (stripped_line_text, original_line_index)

        for original_idx, line_text in enumerate(original_lines):
            stripped = line_text.strip()
            if stripped: # Only consider non-empty lines for parsing logic
                processed_lines_info.append({'text': stripped, 'original_index': original_idx})

        if not processed_lines_info:
            return None, 0

        header_line_proc_index = -1 # Index in processed_lines_info
        separator_line_proc_index = -1
        
        separator_pattern = r"^\s*\|(\s*[:\-]+\s*\|)*\s*[:\-]+\s*\|\s*$"

        # Find header and separator lines using processed_lines_info
        for i, current_line_info in enumerate(processed_lines_info):
            current_line_text = current_line_info['text']
            if not current_line_text.startswith('|') or not current_line_text.endswith('|'):
                continue
            if current_line_text.count('|') < 2: 
                continue

            if (i + 1) < len(processed_lines_info):
                next_line_info = processed_lines_info[i+1]
                next_line_text = next_line_info['text']
                if re.fullmatch(separator_pattern, next_line_text):
                    temp_headers = [h.strip() for h in current_line_text[1:-1].split('|')]
                    num_header_cols = len(temp_headers)
                    num_separator_cols = next_line_text.count('|') - 1
                    
                    if num_header_cols == num_separator_cols and num_header_cols > 0:
                        header_line_proc_index = i
                        separator_line_proc_index = i + 1
                        break
        
        if header_line_proc_index == -1:
            return None, 0

        header_text_from_processed = processed_lines_info[header_line_proc_index]['text']
        headers = [h.strip() for h in header_text_from_processed[1:-1].split('|')]
        num_cols = len(headers)

        table_rows_data = []
        last_processed_row_proc_index = separator_line_proc_index

        # Process rows after the separator line, using processed_lines_info
        for i in range(separator_line_proc_index + 1, len(processed_lines_info)):
            current_row_info = processed_lines_info[i]
            current_row_text = current_row_info['text']

            if current_row_text.startswith('|') and current_row_text.endswith('|'):
                if current_row_text.count('|') != num_cols + 1:
                    break 
                cells = [cell.strip() for cell in current_row_text[1:-1].split('|')]
                if len(cells) == num_cols:
                     table_rows_data.append(cells)
                else: # Fallback for mismatched cell count (should be rare with pipe check)
                    if len(cells) < num_cols: cells.extend([""] * (num_cols - len(cells)))
                    else: cells = cells[:num_cols]
                    table_rows_data.append(cells)
                last_processed_row_proc_index = i # Update index of the last successfully processed row
            else:
                break

        # Determine lines consumed from the original input string
        # The table ends at the original line index of the last processed row (header, separator, or data row)
        # If no rows, table ends at separator. If rows, table ends at last data row.
        # The number of lines consumed is the original_index of the last part of the table + 1
        # (because original_index is 0-based).

        # If table_rows_data is not empty, the table consumed lines up to the last data row.
        # The index in processed_lines_info for this last data row is `last_processed_row_proc_index`.
        # If table_rows_data is empty, the table consumed lines up to the separator line.
        # The index in processed_lines_info for the separator is `separator_line_proc_index`.

        # `last_processed_row_proc_index` is initialized to `separator_line_proc_index`
        # and updated if any data rows are found and processed.
        # So, `processed_lines_info[last_processed_row_proc_index]['original_index']` gives the
        # 0-based index in the *original* `markdown_text.splitlines()` of the last line
        # that is part of the parsed table.
        lines_consumed_count = processed_lines_info[last_processed_row_proc_index]['original_index'] + 1

        return {'type': 'table', 'headers': headers, 'rows': table_rows_data}, lines_consumed_count

    def _parse_implicit_table(self, text_lines: list[str]) -> dict or None:
        MIN_IMPLICIT_TABLE_ROWS = 2 # Minimum number of qualifying lines to form a table
        # Regex to capture key and value parts. Allows for optional whitespace around colon.
        # Key is group 1, Value is group 2.
        # Key can be anything up to the colon. Value is everything after.
        KEY_VALUE_SEPARATOR_PATTERN = re.compile(r"^\s*(.+?)\s*:\s*(.+)\s*$")

        if not text_lines or len(text_lines) < MIN_IMPLICIT_TABLE_ROWS:
            return None

        parsed_rows_data = []
        expected_num_value_cols = -1 # Undetermined initially

        # Find the first valid line to determine column count and start the table
        start_line_idx = -1
        first_row_candidate = []

        for idx, line in enumerate(text_lines):
            line_stripped = line.strip()
            if not line_stripped: # Skip blank lines between potential data lines
                if start_line_idx != -1 and not parsed_rows_data: # Blank line before any valid rows collected after start
                    start_line_idx = -1 # Reset start if a blank line interrupts before min rows
                continue

            match = KEY_VALUE_SEPARATOR_PATTERN.match(line_stripped)
            if match:
                key_part = match.group(1).strip()
                value_part = match.group(2).strip()

                # Split value_part by comma, trim spaces from each cell
                value_cells = [v.strip() for v in value_part.split(',')]

                if not value_cells or not value_cells[0]: # Value part must exist
                    if start_line_idx != -1: # If we were already building a table
                        break # Invalid line, stop table parsing here
                    else: continue # Keep searching for a valid start line

                current_row_values = [key_part] + value_cells
                num_value_cols_this_line = len(value_cells)

                if start_line_idx == -1: # This is the first valid line candidate
                    start_line_idx = idx
                    expected_num_value_cols = num_value_cols_this_line
                    parsed_rows_data.append(current_row_values)
                elif num_value_cols_this_line == expected_num_value_cols:
                    # This line matches the expected column structure
                    parsed_rows_data.append(current_row_values)
                else:
                    # Line does not match column structure, table ends here
                    break
            else:
                # Line does not match "Key: Value" pattern.
                if start_line_idx != -1: # If we were already building a table
                    break # Non-matching line, stop table parsing here
                # else: continue searching for a start line (handles leading non-matching lines)

        if len(parsed_rows_data) >= MIN_IMPLICIT_TABLE_ROWS:
            # Construct headers
            headers = ["Parameter"]
            for i in range(expected_num_value_cols):
                # Using "Value {i+1}" but could be "Comp {i+1}" etc. if context allows
                headers.append(f"Value {i+1}")

            # The number of lines consumed by this implicit table is tricky if there were interspersed blank lines
            # or leading/trailing non-matching lines. For now, this function doesn't return consumed lines.
            # _format_ai_response will have to manage line consumption more carefully if this parser is used.
            return {'type': 'table', 'headers': headers, 'rows': parsed_rows_data}

        return None

    def _populate_comparison_treeview(self, ai_response_text: str):
        if hasattr(self, 'comparison_treeview'):
            for item in self.comparison_treeview.get_children(): self.comparison_treeview.delete(item)
        else: self.update_conversation_history("System: Treeview not found.", role="error"); return
        
        # Attempt to parse the response as a generic markdown table first
        # MODIFIED to handle tuple return from _parse_markdown_table
        parsing_result = self._parse_markdown_table(ai_response_text)
        parsed_table_data = None # Initialize
        lines_consumed = 0 # Initialize

        if isinstance(parsing_result, tuple) and len(parsing_result) == 2:
            parsed_table_data, lines_consumed = parsing_result # Unpack tuple
        elif isinstance(parsing_result, dict):
            # Fallback: if _parse_markdown_table somehow returned only a dict (older version or specific path)
            parsed_table_data = parsing_result
            print("DEBUG: _populate_comparison_treeview - _parse_markdown_table returned a dict directly.")
        elif parsing_result is None:
            # _parse_markdown_table returned None directly (e.g. if input was empty, or no table found which now returns (None,0))
            # This specific None case might be less likely now with (None,0) return for "no table found"
            print("DEBUG: _populate_comparison_treeview - _parse_markdown_table returned None directly.")
        else:
            # Unexpected return type
            print(f"DEBUG: _populate_comparison_treeview - Unexpected return type from _parse_markdown_table: {type(parsing_result)}")

        # Ensure parsed_table_data is a dict and has 'headers' and 'rows' keys before proceeding
        if not parsed_table_data or not isinstance(parsed_table_data, dict) or \
           not parsed_table_data.get('headers') or not parsed_table_data.get('rows'):
            self.update_conversation_history("System: No valid table data parsed for Treeview or table is empty/malformed.", role="system")
            return

        pn1 = self.mfg_pn_var_1.get() or (os.path.basename(self.spec_sheet_1_path) if self.spec_sheet_1_path else "Comp 1")
        pn2 = self.mfg_pn_var_2.get() or (os.path.basename(self.spec_sheet_2_path) if self.spec_sheet_2_path else "Comp 2")
        
        headers = parsed_table_data.get('headers', []) # Use .get for safety
        
        self.comparison_treeview.heading("component1", text=f"{pn1[:25]}{'...' if len(pn1)>25 else ''}")
        self.comparison_treeview.heading("component2", text=f"{pn2[:25]}{'...' if len(pn2)>25 else ''}")

        for row_data in parsed_table_data.get('rows', []): # Use .get for safety
            # Map row_data list to the tuple expected by treeview.insert
            # (parameter, component1_val, component2_val, notes)
            parameter = row_data[0] if len(row_data) > 0 else ""
            comp1_val = row_data[1] if len(row_data) > 1 else ""
            comp2_val = row_data[2] if len(row_data) > 2 else ""
            notes = row_data[3] if len(row_data) > 3 else ""
            self.comparison_treeview.insert("", tk.END, values=(parameter, comp1_val, comp2_val, notes))
            
        self.update_conversation_history("System: Detailed comparison table populated.", role="system")

    def load_spec_sheet_1(self):
        filepath = filedialog.askopenfilename(title="Select Spec Sheet 1 (PDF)", filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")])
        if filepath:
            self.spec_sheet_1_path = filepath
            self.spec_sheet_1_label.config(text=f"File 1: {os.path.basename(filepath)}")
            if self.spec_sheet_1_path and self.spec_sheet_2_path:
                self.model_combobox.config(state='readonly')
                self.update_conversation_history("System: Both spec sheets loaded. Please select an AI model.", role="system")
        else:
            self.spec_sheet_1_label.config(text=f"File 1: {os.path.basename(self.spec_sheet_1_path) if self.spec_sheet_1_path else 'None'}")

    def load_spec_sheet_2(self):
        filepath = filedialog.askopenfilename(title="Select Spec Sheet 2 (PDF)", filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")])
        if filepath:
            self.spec_sheet_2_path = filepath
            self.spec_sheet_2_label.config(text=f"File 2: {os.path.basename(filepath)}")
            if self.spec_sheet_1_path and self.spec_sheet_2_path:
                self.model_combobox.config(state='readonly')
                self.update_conversation_history("System: Both spec sheets loaded. Please select an AI model.", role="system")
        else:
            self.spec_sheet_2_label.config(text=f"File 2: {os.path.basename(self.spec_sheet_2_path) if self.spec_sheet_2_path else 'None'}")

    
    def _is_text_segment_redundant_with_table(self, text_lines: list[str], table_data: dict) -> bool:
        MAX_LINES_FOR_REDUNDANCY_CHECK = 7
        MIN_REDUNDANCY_THRESHOLD_PERCENT = 0.75
        MIN_REDUNDANCY_FOR_SECTION_MATCH = 0.51

        actual_text_lines = [line for line in text_lines if line.strip()]
        if not actual_text_lines:
            return False

        is_potentially_headed_section = False
        first_line_normalized_for_header_check = actual_text_lines[0].strip().lower()
        if (first_line_normalized_for_header_check.startswith('**') and first_line_normalized_for_header_check.endswith('**') and len(first_line_normalized_for_header_check) > 4) or        (first_line_normalized_for_header_check.startswith('## ') and len(first_line_normalized_for_header_check) > 3):
            is_potentially_headed_section = True

        if not is_potentially_headed_section and len(actual_text_lines) > MAX_LINES_FOR_REDUNDANCY_CHECK:
            return False

        table_headers_normalized = {str(h).strip().lower(): str(h).strip() for h in table_data.get('headers', [])}
        table_all_cells_normalized_set = set()
        for row in table_data.get('rows', []):
            for cell in row:
                table_all_cells_normalized_set.add(str(cell).strip().lower())

        if not table_all_cells_normalized_set and not table_headers_normalized:
            return False

        potential_section_title_normalized = None
        first_line_orig_case = actual_text_lines[0].strip()

        if first_line_orig_case.startswith('**') and first_line_orig_case.endswith('**') and len(first_line_orig_case) > 4:
            potential_section_title_normalized = first_line_orig_case[2:-2].strip().lower()
        elif first_line_orig_case.startswith('## ') and len(first_line_orig_case) > 3:
            potential_section_title_normalized = first_line_orig_case[3:].strip().lower()

        lines_to_check_for_content = actual_text_lines
        focused_table_content_normalized = None
        is_section_header_matched = False

        if potential_section_title_normalized:
            if potential_section_title_normalized in table_headers_normalized:
                is_section_header_matched = True
                try:
                    col_idx = -1
                    for idx, header_val in enumerate(table_data.get('headers', [])):
                        if str(header_val).strip().lower() == potential_section_title_normalized:
                            col_idx = idx
                            break
                    if col_idx != -1:
                        focused_table_content_normalized = {str(row[col_idx]).strip().lower() for row in table_data.get('rows', []) if len(row) > col_idx and str(row[col_idx]).strip()}
                except Exception:
                    pass
                lines_to_check_for_content = actual_text_lines[1:]
            else:
                if table_data.get('rows') and len(table_data['rows'][0]) > 0:
                    for row_data in table_data.get('rows', []):
                        if str(row_data[0]).strip().lower() == potential_section_title_normalized:
                            is_section_header_matched = True
                            focused_table_content_normalized = {str(cell).strip().lower() for cell in row_data[1:] if str(cell).strip()}
                            lines_to_check_for_content = actual_text_lines[1:]
                            break

        if is_section_header_matched and not [line for line in lines_to_check_for_content if line.strip()]:
            return True # Header matched, no content lines, considered redundant

        if len([line for line in lines_to_check_for_content if line.strip()]) > MAX_LINES_FOR_REDUNDANCY_CHECK and not is_section_header_matched:
             return False

        redundant_lines_count = 0
        comparison_basis_set = focused_table_content_normalized if focused_table_content_normalized is not None else table_all_cells_normalized_set

        if not comparison_basis_set:
            if is_section_header_matched :
                 return False

        content_lines_for_final_check = [line for line in lines_to_check_for_content if line.strip()][:MAX_LINES_FOR_REDUNDANCY_CHECK]
        if not content_lines_for_final_check:
            return False

        for line_content in content_lines_for_final_check:
            normalized_line = line_content.strip().lower()
            if normalized_line in comparison_basis_set:
                redundant_lines_count += 1; continue

            found_substring_match = False
            for table_val in comparison_basis_set:
                if table_val and (table_val in normalized_line or normalized_line in table_val):
                    found_substring_match = True; break
            if found_substring_match:
                redundant_lines_count += 1; continue

            if focused_table_content_normalized is None and ':' in normalized_line: # Only do key:value check for general text
                parts = normalized_line.split(':', 1)
                if len(parts) == 2:
                    key_norm = parts[0].strip()
                    val_norm = parts[1].strip()
                    # Check if key is a table header and value is in any cell, or both in general cells
                    if (key_norm in table_headers_normalized and val_norm in table_all_cells_normalized_set) or                    (key_norm in table_all_cells_normalized_set and val_norm in table_all_cells_normalized_set):
                        redundant_lines_count += 1; continue

        current_threshold = MIN_REDUNDANCY_FOR_SECTION_MATCH if is_section_header_matched and focused_table_content_normalized is not None else MIN_REDUNDANCY_THRESHOLD_PERCENT

        if (float(redundant_lines_count) / len(content_lines_for_final_check)) >= current_threshold:
            return True

        return False

    def _finalize_text_block(self, block_lines: list[str], last_table_segment_for_redundancy_check: dict or None) -> dict or None:
        if not [line for line in block_lines if line.strip()]:
            return None

        implicit_table = self._parse_implicit_table(block_lines)
        if implicit_table:
            # REMOVED: print(f"DEBUG: _finalize_text_block: Added IMPLICIT TABLE. Headers: {{implicit_table.get('headers')}}")
            return implicit_table

        collected_text = "\n".join(block_lines).strip()
        if collected_text:
            is_redundant = False
            if last_table_segment_for_redundancy_check:
                is_redundant = self._is_text_segment_redundant_with_table(block_lines, last_table_segment_for_redundancy_check)

            if not is_redundant:
                # REMOVED: print(f"DEBUG: _finalize_text_block: Added TEXT segment: '{{collected_text[:70]}}...'")
                return {'type': 'text', 'content': collected_text}
            else:
                # REMOVED: print(f"DEBUG: _finalize_text_block: Suppressed redundant TEXT segment: '{{collected_text[:70]}}...'")
                pass # Explicitly do nothing if text is redundant and suppressed
        return None

    def _format_ai_response(self, text_response: str) -> list:
        segments = []
        current_text_block_lines = [] # Accumulates lines for a potential text or implicit table segment
        all_lines = text_response.splitlines()
        i = 0
        last_processed_table_segment = None # For redundancy checks of text vs preceding table

        while i < len(all_lines):
            # Check for a standard markdown (pipe) table starting at the current line
            # _parse_markdown_table expects a single string, so we join lines from current position
            pipe_table_data, lines_consumed_by_parser = self._parse_markdown_table("\n".join(all_lines[i:]))

            if pipe_table_data:
                # A pipe table was found. First, finalize any text block accumulated *before* this pipe table.
                if current_text_block_lines:
                    processed_segment = self._finalize_text_block(current_text_block_lines, last_processed_table_segment)
                    if processed_segment:
                        segments.append(processed_segment)
                        if processed_segment['type'] == 'table': # This would be an implicit table
                            last_processed_table_segment = processed_segment
                    current_text_block_lines = [] # Reset buffer

                # Now, add the pipe table itself
                segments.append(pipe_table_data)
                last_processed_table_segment = pipe_table_data # Update for next redundancy checks
                
                # Use the accurate lines_consumed_by_parser from the parsing method
                # REMOVED: print(f"DEBUG: _format_ai_response: Added PIPE TABLE segment. Headers: {pipe_table_data.get('headers')}, Consumed: {lines_consumed_by_parser} lines")
                i += lines_consumed_by_parser # Advance 'i' by the number of lines consumed by the table parser

            else: # No pipe table starts at all_lines[i]
                line = all_lines[i]
                if not line.strip(): # Current line is blank, signifies end of a text block
                    if current_text_block_lines:
                        processed_segment = self._finalize_text_block(current_text_block_lines, last_processed_table_segment)
                        if processed_segment:
                            segments.append(processed_segment)
                            if processed_segment['type'] == 'table': # Implicit table
                                last_processed_table_segment = processed_segment
                        current_text_block_lines = [] # Reset buffer
                else: # Non-blank line, add to current text block
                    current_text_block_lines.append(line)
                i += 1 # Move to the next line

        # After the loop, finalize any remaining lines in current_text_block_lines
        if current_text_block_lines:
            processed_segment = self._finalize_text_block(current_text_block_lines, last_processed_table_segment)
            if processed_segment:
                segments.append(processed_segment)
                # No need to update last_processed_table_segment here as it's the end.

        # Debug print for the final list of segments
        # REMOVED: print(f"DEBUG: _format_ai_response: RETURNING segments (count {len(segments)}): {[s['type'] for s in segments]}") # Log segment types
        return segments

    def clean_cell_content(cell_text):
            # Remove leading/trailing whitespace and markdown bold markers
            return cell_text.strip().replace("**", "")

    def format_table(table_lines):
        if not table_lines:
            return ""

        parsed_table_inner = [] # Renamed
        for line_inner in table_lines: # Renamed
            # Remove leading/trailing '|' and split by '|'
            cells = [clean_cell_content(cell) for cell in line_inner.strip()[1:-1].split('|')]
            parsed_table_inner.append(cells)

        if not parsed_table_inner:
            return "\n".join(table_lines) # Should not happen if table_lines is not empty

        num_cols = len(parsed_table_inner[0])
        # Ensure all rows have the same number of columns, pad if necessary
        for row_idx, row_inner in enumerate(parsed_table_inner): # Renamed
            if len(row_inner) < num_cols:
                parsed_table_inner[row_idx].extend([""] * (num_cols - len(row_inner)))
            elif len(row_inner) > num_cols: # Should ideally not happen with well-formed tables
                parsed_table_inner[row_idx] = row_inner[:num_cols]


        col_widths = [0] * num_cols
        for row_inner_widths in parsed_table_inner: # Renamed
            for i_cell, cell_content in enumerate(row_inner_widths): # Renamed
                if i_cell < num_cols: # Ensure we don't go out of bounds
                    col_widths[i_cell] = max(col_widths[i_cell], len(cell_content))

        formatted_table_str_lines = []
        for i_format_row, row_format_data in enumerate(parsed_table_inner): # Renamed
            formatted_row_parts = []
            for j_format_cell, cell_format_data in enumerate(row_format_data): # Renamed
                if j_format_cell < num_cols: # Ensure we don't go out of bounds
                    # For separator row, create the separator line based on calculated widths
                    if separator_pattern_local.match(table_lines[i_format_row].strip()): # Check original line for separator
                        formatted_row_parts.append('-' * col_widths[j_format_cell])
                    else:
                        formatted_row_parts.append(cell_format_data.ljust(col_widths[j_format_cell]))
            formatted_table_str_lines.append(" | ".join(formatted_row_parts))

        # Re-add outer pipes for aesthetics if desired, or leave as is for simpler alignment
        return "\n".join([f"| {s} |" for s in formatted_table_str_lines])


        for i_outer, line_outer in enumerate(lines): # Renamed
            stripped_line = line_outer.strip()

            is_table_row_candidate = table_row_pattern.match(stripped_line)
            is_separator = separator_pattern_local.match(stripped_line)

            if is_table_row_candidate:
                if not in_table_block:
                    # Check if this is the start of a new table
                    # Look ahead for a separator or more table rows
                    if is_separator or \
                       (i_outer + 1 < len(lines) and (separator_pattern_local.match(lines[i_outer+1].strip()) or table_row_pattern.match(lines[i_outer+1].strip()))):
                        in_table_block = True
                        current_table_lines.append(stripped_line)
                    else: # Not a table, just a line with pipes
                        formatted_output_lines.append(re.sub(r'\s*\|\s*', ' | ', stripped_line))
                else: # Already in a table block
                    current_table_lines.append(stripped_line)
            else: # Not a table row candidate
                if in_table_block:
                    # Table block has ended
                    formatted_output_lines.append(format_table(current_table_lines))
                    current_table_lines = []
                    in_table_block = False
                # Add the non-table line (could be empty)
                formatted_output_lines.append(line_outer) # Keep original non-table lines

        # If the response ends with a table block
        if in_table_block and current_table_lines:
            formatted_output_lines.append(format_table(current_table_lines))

        return "\n".join(formatted_output_lines) # This also needs to change

    def update_conversation_history(self, message, role="system"):
        raw_message_for_log = message # Keep the original message for the log

        if hasattr(self, 'conversation_history') and self.conversation_history:
            self.conversation_history.config(state=tk.NORMAL)
            tag_to_apply = {"user": "user_message", "ai": "ai_message", "error": "error_message"}.get(role, "system_message")

            if role == "ai":
                formatted_content_or_segments = self._format_ai_response(raw_message_for_log) # Use raw message for formatting
                
                # This part needs to handle a list of segments
                if isinstance(formatted_content_or_segments, list):
                    for segment in formatted_content_or_segments:
                        if segment['type'] == 'table':
                            headers = segment.get('headers', [])
                            rows = segment.get('rows', [])
                            if headers: # Only proceed if there's actual table data
                                table_frame = ttk.Frame(self.conversation_history)
                                column_ids = [f"col_{i}" for i, _ in enumerate(headers)]
                                tree = ttk.Treeview(table_frame, columns=column_ids, show="headings", height=len(rows) if rows else 1)
                                for i, header_text in enumerate(headers):
                                    tree.heading(column_ids[i], text=header_text.strip(), anchor=tk.W)
                                    tree.column(column_ids[i], anchor=tk.W, width=100, stretch=tk.YES)
                                for row_data in rows:
                                    processed_row = []
                                    for cell_idx, cell in enumerate(row_data):
                                        if cell_idx < len(column_ids):
                                            cell_text = str(cell).replace('\n', ' ').replace('<br>', ' ')
                                            processed_row.append(cell_text)
                                    tree.insert("", tk.END, values=processed_row)
                                tree.pack(side=tk.LEFT, fill=tk.X, expand=True)
                                self.conversation_history.insert(tk.END, '\n', tag_to_apply)
                                self.conversation_history.window_create(tk.END, window=table_frame)
                                self.conversation_history.insert(tk.END, '\n', tag_to_apply)
                            else: # Empty table or malformed
                                self.conversation_history.insert(tk.END, "AI Table (empty or malformed)\n", tag_to_apply)
                        elif segment['type'] == 'text':
                            self.conversation_history.insert(tk.END, segment['content'] + "\n", tag_to_apply)
                elif isinstance(formatted_content_or_segments, dict) and formatted_content_or_segments.get('type') == 'table':
                    # This case handles the old behavior where _format_ai_response might directly return one table
                    # This should ideally be deprecated by the new list-based approach
                    headers = formatted_content_or_segments.get('headers', [])
                    rows = formatted_content_or_segments.get('rows', [])
                    if headers: 
                        table_frame = ttk.Frame(self.conversation_history)
                        column_ids = [f"col_{i}" for i, _ in enumerate(headers)]
                        tree = ttk.Treeview(table_frame, columns=column_ids, show="headings", height=len(rows) if rows else 1)
                        for i, header_text in enumerate(headers):
                            tree.heading(column_ids[i], text=header_text.strip(), anchor=tk.W)
                            tree.column(column_ids[i], anchor=tk.W, width=100, stretch=tk.YES)
                        for row_data in rows:
                            processed_row = []
                            for cell_idx, cell in enumerate(row_data):
                                if cell_idx < len(column_ids):
                                    cell_text = str(cell).replace('\n', ' ').replace('<br>', ' ')
                                    processed_row.append(cell_text)
                            tree.insert("", tk.END, values=processed_row)
                        tree.pack(side=tk.LEFT, fill=tk.X, expand=True)
                        self.conversation_history.insert(tk.END, '\n', tag_to_apply)
                        self.conversation_history.window_create(tk.END, window=table_frame)
                        self.conversation_history.insert(tk.END, '\n', tag_to_apply)
                    else:
                        self.conversation_history.insert(tk.END, "AI Table (empty or malformed)\n", tag_to_apply)
                else: # It's formatted text (string) - old fallback
                    self.conversation_history.insert(tk.END, str(formatted_content_or_segments) + "\n", tag_to_apply)
            else: # User, system, error messages
                display_message = raw_message_for_log
                self.conversation_history.insert(tk.END, display_message + "\n", tag_to_apply)

            self.conversation_history.see(tk.END)
            self.conversation_history.config(state=tk.DISABLED)
        
        self.conversation_log.append({'role': role, 'content': raw_message_for_log})

    def _update_ui_for_ai_status(self, api_key_configured=None, model_initialized=None):
        if not hasattr(self, 'send_button'): return
        is_api_key_ready = api_key_configured if api_key_configured is not None else self.api_key_configured
        is_model_ready = model_initialized if model_initialized is not None else (self.model is not None)
        can_perform_ai_ops = is_api_key_ready and is_model_ready
        self.send_button.config(state=tk.NORMAL if can_perform_ai_ops else tk.DISABLED)
        self.user_input_entry.config(state=tk.NORMAL if can_perform_ai_ops else tk.DISABLED)
        if hasattr(self, 'model_combobox') and self.model_combobox.cget('state') != 'disabled':
            self.model_combobox.config(state="readonly")

    def _configure_ai(self):
        try:
            api_key = os.environ.get("GOOGLE_API_KEY")
            if not api_key: print("DEBUG: GOOGLE_API_KEY not found for _configure_ai."); self.api_key_configured = False
            else: genai.configure(api_key=api_key); self.api_key_configured = True
        except Exception as e: self.update_conversation_history(f"System: Error configuring AI SDK: {e}", role="error"); self.api_key_configured = False
        finally: self._update_ui_for_ai_status(api_key_configured=self.api_key_configured, model_initialized=(self.model is not None))

    def _on_model_selected(self, event=None):
        if event: print(f"DEBUG: _on_model_selected. Event: {event.type}, Widget: {event.widget}")
        else: print(f"DEBUG: _on_model_selected programmatically.")
        selected_model_name = self.model_combobox.get(); print(f"DEBUG: Combobox get(): '{selected_model_name}'")
        if selected_model_name == self.placeholder_text:
            self.update_conversation_history("System: Select a valid AI model.", role="system"); self._update_ui_for_ai_status(model_initialized=False); return
        previous_model_name = self.model.model_name if self.model else None
        is_diff_model = self.model and self.model.model_name != selected_model_name

        is_first_select_with_history = not self.model and self.conversation_log and \
            any(not (log.get('role') == 'system' and log.get('content', '').startswith("System: Welcome!")) for log in self.conversation_log)

        if is_diff_model or is_first_select_with_history:
            log_msg = f"System: Changing model";
            if previous_model_name: log_msg += f" from {previous_model_name}"
            log_msg += f" to {selected_model_name}. Clearing context."
            self.update_conversation_history(log_msg, role="system"); self.clear_all(clear_files=False)
            self.update_conversation_history(f"System: AI Model selected: {selected_model_name}", role="system")
        else: self.update_conversation_history(f"System: AI Model selected: {selected_model_name}", role="system")
        self.model_initializing = True
        model_init_ok = self._initialize_model(selected_model_name)
        self.model_initializing = False
        if model_init_ok:
            if self.spec_sheet_1_path and self.spec_sheet_2_path:
                self.update_conversation_history(f"System: Starting analysis with {selected_model_name}...", role="system")
                self.check_and_process_spec_sheets()
            else: self.update_conversation_history("System: Model initialized. Load both spec sheets.", role="system")
        else: self.update_conversation_history(f"System: Failed to init {selected_model_name}. Check logs.", role="error")

    def _initialize_model(self, model_name=None):
        source_log = "explicitly";
        if model_name is None: model_name = self.model_combobox.get(); source_log = f"from Combobox: '{model_name}'"
        if model_name == self.placeholder_text:
            self.update_conversation_history("System: Select valid model.", role="system"); self.model=None; self.chat_session=None; self._update_ui_for_ai_status(model_initialized=False); return False
        self.update_conversation_history(f"System: _initialize_model ({source_log}). Target: '{model_name}'", role="system")
        if not model_name or model_name not in self.model_options_list:
            self.update_conversation_history(f"System: Invalid model ('{model_name}'). Aborted.", role="error"); self.model=None; self.chat_session=None; self._update_ui_for_ai_status(model_initialized=False); return False
        if self.model and self.model.model_name == model_name and self.api_key_configured:
            self.update_conversation_history(f"System: Model '{model_name}' already active.", role="system"); self._update_ui_for_ai_status(model_initialized=True); return True
        if not self.api_key_configured:
            self.update_conversation_history("System: Cannot init model - API key not set.", role="error"); self.model=None; self.chat_session=None; self._update_ui_for_ai_status(model_initialized=False); return False
        self.update_conversation_history(f"System: Initializing model: {model_name}...", role="system")
        try:
            if not any(isinstance(log, dict) and "Generative AI configured successfully." in log.get('content', '') for log in self.conversation_log):
                self.update_conversation_history("System: Generative AI configured successfully.", role="system")
            self.model = genai.GenerativeModel(model_name); self.chat_session = None
            self.update_conversation_history(f"System: Successfully initialized model: {model_name}", role="system"); self._update_ui_for_ai_status(model_initialized=True); return True
        except Exception as e: self.model=None; self.chat_session=None; self.update_conversation_history(f"System: Error initializing model {model_name}: {e}", role="error"); self._update_ui_for_ai_status(model_initialized=False); return False

    def get_selected_model_name(self): return self.model_var.get()
    def _add_to_ai_history(self,role:str,text_content:str): self.ai_history.append({'role':role,'parts':[text_content]}); print(f"DEBUG: AI history add: {role}, '{text_content[:50]}...'")
    def _convert_log_to_gemini_history(self): return [e for e in self.ai_history if e['role'] in ('user','model')]

    def send_user_query(self):
        if not self.model or not self.api_key_configured:
            if not self._initialize_model(): return
        if not self.model: self.update_conversation_history("System: AI Model N/A.", role="error"); return

        user_text = self.user_input_entry.get().strip()
        # self.update_conversation_history is called after constructing prompt_parts_for_ai
        self.user_input_entry.delete(0, tk.END)

        prompt_parts_for_ai = []
        image_sent_this_turn = False
        log_message_for_user_turn = user_text

        if self.pending_user_image_pil:
            if user_text: # Text and image
                prompt_parts_for_ai.append(user_text)
                prompt_parts_for_ai.append(self.pending_user_image_pil)
            else: # Image only
                prompt_parts_for_ai.append(self.pending_user_image_pil)

            image_filename = os.path.basename(self.pending_user_image_path)
            self.update_conversation_history(f"User: {user_text} [Image: {image_filename}]", role="user")
            log_message_for_user_turn = f"{user_text} [Image: {image_filename}]"
            image_sent_this_turn = True
        elif user_text: # Text only
            prompt_parts_for_ai.append(user_text)
            self.update_conversation_history(f"User: {user_text}", role="user")
        else:
            self.update_conversation_history("System: Cannot send empty message.", role="system"); return

        active_model_name = self.model.model_name
        try:
            self.send_button.config(state=tk.DISABLED); self.user_input_entry.config(state=tk.DISABLED)
            if not self.chat_session:
                self.update_conversation_history(f"System: Starting new chat with {active_model_name}...", role="system")
                try: self.chat_session = self.model.start_chat(history=self._convert_log_to_gemini_history())
                except Exception as e: self.update_conversation_history(f"System: Error starting chat: {e}", role="error"); self._update_ui_for_ai_status(); return

            self._add_to_ai_history('user', log_message_for_user_turn)

            self.update_conversation_history(f"System: Sending to AI ({active_model_name})...", role="system")

            datasheet_sourcing_instruction_text = (
                "IMPORTANT INSTRUCTIONS FOR AI RESPONSE:\n"
                "- Base your answers primarily on the information extracted from the provided component datasheets "
                "(text, images, and context from previous conversation turns, including component type and MFG P/N if known).\n"
                "- If the datasheets or prior conversation context lack specific information to answer your query, you may use your general knowledge.\n"
                "- If you use general knowledge, you MUST explicitly state that the information was not found in the provided datasheets/context "
                "and that your answer is based on general understanding.\n\n"
            )

            # Construct the prompt with sourcing instruction first, then user's actual content parts
            prompt_with_sourcing_instruction = [datasheet_sourcing_instruction_text] + prompt_parts_for_ai

            final_prompt_parts_for_sending = list(prompt_with_sourcing_instruction) # Use a copy

            if hasattr(self, 'root'): self.root.update_idletasks()
            if self.translate_to_chinese_var.get():
                translation_instruction = " Please provide your entire response in Chinese."
            else:
                translation_instruction = " Please provide your entire response in English."

            appended_to_text = False
            for i in range(len(final_prompt_parts_for_sending) - 1, -1, -1):
                if isinstance(final_prompt_parts_for_sending[i], str):
                    final_prompt_parts_for_sending[i] += translation_instruction
                    appended_to_text = True; break
            if not appended_to_text: final_prompt_parts_for_sending.append(translation_instruction)

            response = self.chat_session.send_message(final_prompt_parts_for_sending)
            self._add_to_ai_history('model', response.text)
            self.update_conversation_history(f"AI ({active_model_name}): {response.text}", role="ai")

            if image_sent_this_turn:
                self.pending_user_image_path = None; self.pending_user_image_pil = None
        except Exception as e:
            err_msg=f"System: Error with AI ({active_model_name}): {e}"; self.update_conversation_history(err_msg,role="error"); print(f"DEBUG: {err_msg}")
            if isinstance(e,(google_exceptions.PermissionDenied,google_exceptions.Unauthenticated)): self.api_key_configured=False
            if isinstance(e,(google_exceptions.InvalidArgument,ValueError,BlockedPromptException,StopCandidateException,google_exceptions.NotFound,google_exceptions.PermissionDenied)):
                self.update_conversation_history(f"System: Resetting model ({active_model_name}) due to error.",role="system"); self.model=None; self.chat_session=None; self.ai_history=[]
        finally: self._update_ui_for_ai_status()

    def download_history(self):
        # Ensure necessary imports are at the top of main.py:
        # import os
        # import docx
        # from tkinter import filedialog # Already imported usually
        # from docx.shared import RGBColor # Should be at top of main.py by now
        # import traceback # Should be at top of main.py by now

        if not self.conversation_log:
            self.update_conversation_history("System: History empty. Nothing to download.", role="system")
            return

        try:
            filepath = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")],
                title="Save Conversation History"
            )
        except Exception as e:
            self.update_conversation_history(f"System: Error opening save dialog: {e}", role="error")
            return

        if not filepath:
            self.update_conversation_history("System: Download cancelled by user.", role="system")
            return

        try:
            doc = docx.Document()
            # RGBColor should be imported via `from docx.shared import RGBColor` at module level

            doc.add_heading("Component Comparator AI Chat History", level=1)

            USER_COLOR = RGBColor(0x00, 0x00, 0xFF)  # Blue
            AI_COLOR = RGBColor(0x00, 0x80, 0x00)    # Green
            SYSTEM_COLOR = RGBColor(0x80, 0x00, 0x80) # Purple
            ERROR_COLOR = RGBColor(0xFF, 0x00, 0x00)   # Red
            DEFAULT_COLOR = RGBColor(0x00, 0x00, 0x00) # Black

            color_map = {
                'user': USER_COLOR,
                'ai': AI_COLOR,
                'system': SYSTEM_COLOR,
                'error': ERROR_COLOR
            }

            if self.spec_sheet_1_path:
                p = doc.add_paragraph()
                # os.path.basename needs `import os` at module level
                run = p.add_run(f"Spec Sheet 1: {os.path.basename(self.spec_sheet_1_path)}")
                run.font.color.rgb = SYSTEM_COLOR
            if self.spec_sheet_2_path:
                p = doc.add_paragraph()
                run = p.add_run(f"Spec Sheet 2: {os.path.basename(self.spec_sheet_2_path)}")
                run.font.color.rgb = SYSTEM_COLOR

            model_name_to_log = "N/A"
            if self.model and hasattr(self.model, 'model_name'):
                model_name_to_log = self.model.model_name
            p = doc.add_paragraph()
            run = p.add_run(f"AI Model (last used): {model_name_to_log}")
            run.font.color.rgb = SYSTEM_COLOR
            doc.add_paragraph("-" * 20)

            for entry_data in self.conversation_log:
                # CRUCIAL DEBUG LOGGING TO VERIFY THIS VERSION IS RUNNING:
                # print(f"DEBUG download_history (v3): Processing entry_data of type: {type(entry_data)}, content snippet: '{str(entry_data)[:150]}...'")

                entry_role = None
                entry_content = None

                if isinstance(entry_data, dict):
                    entry_role = entry_data.get('role', 'system')
                    entry_content = entry_data.get('content', '')
                elif isinstance(entry_data, str):
                    entry_role = 'system'
                    entry_content = entry_data
                    # REMOVED: print(f"DEBUG download_history (v3): Handled string entry: '{entry_data[:100]}...'. Assigned role '{entry_role}'.")
                else:
                    # REMOVED: print(f"DEBUG download_history (v3): Skipping unknown entry type in log: {type(entry_data)}")
                    continue

                if entry_content is None:
                    entry_content = ''

                if not isinstance(entry_content, str):
                    entry_content = str(entry_content)

                text_color = color_map.get(entry_role, DEFAULT_COLOR)

                if not entry_content.strip():
                    continue

                segments = self._format_ai_response(entry_content)

                if segments:
                    for segment_idx, segment in enumerate(segments):
                        segment_type = segment.get('type')

                        if segment_type == 'table':
                            headers = segment.get('headers', [])
                            data_rows = segment.get('rows', [])
                            num_cols = len(headers)

                            if num_cols > 0 and data_rows:
                                word_table = doc.add_table(rows=1, cols=num_cols)
                                word_table.style = 'TableGrid'
                                for col_idx, header_text_val in enumerate(headers):
                                    cell_run = word_table.cell(0, col_idx).paragraphs[0].add_run(str(header_text_val))
                                    cell_run.font.color.rgb = text_color
                                for data_row_list in data_rows:
                                    row_cells = word_table.add_row().cells
                                    for col_idx, cell_text_content in enumerate(data_row_list):
                                        if col_idx < num_cols:
                                            cell_run = row_cells[col_idx].paragraphs[0].add_run(str(cell_text_content))
                                            cell_run.font.color.rgb = text_color
                                if segment_idx < len(segments) - 1:
                                    doc.add_paragraph('')
                            elif num_cols > 0 and not data_rows:
                                p = doc.add_paragraph()
                                run = p.add_run(f"[Table with headers: {', '.join(headers)} - No data rows]")
                                run.font.color.rgb = text_color
                                if segment_idx < len(segments) - 1:
                                    doc.add_paragraph('')

                        elif segment_type == 'text':
                            current_text_content = segment.get('content', '')
                            if current_text_content.strip():
                                p = doc.add_paragraph()
                                run = p.add_run(current_text_content)
                                run.font.color.rgb = text_color
                                if segment_idx < len(segments) - 1:
                                    doc.add_paragraph('')

                elif entry_content.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(f"[Unprocessed Entry - Role: {entry_role}]: {entry_content}")
                    run.font.color.rgb = text_color
                    # REMOVED: print(f"DEBUG: Download History (v3) - Entry was not processed into segments by _format_ai_response (Role: {entry_role}): {entry_content[:100]}...")

            doc.save(filepath)
            self.update_conversation_history(f"System: Conversation history downloaded to {filepath}", role="system")

        except Exception as e:
            # KEPT: This is important error logging
            print(f"DEBUG: Error saving .docx history (v3): {e}\n{traceback.format_exc()}")
            error_message = f"System: Error during history download: {e}"
            self.update_conversation_history(error_message, role="error")

    def clear_all(self, clear_files=True):
        print(f"DEBUG: clear_all called with clear_files={clear_files}")
        if clear_files:
            self.spec_sheet_1_path=None; self.spec_sheet_1_text=None; self.spec_sheet_1_image_paths=[]
            self.spec_sheet_2_path=None; self.spec_sheet_2_text=None; self.spec_sheet_2_image_paths=[]
            if hasattr(self,'spec_sheet_1_label'): self.spec_sheet_1_label.config(text="File 1: None")
            if hasattr(self,'spec_sheet_2_label'): self.spec_sheet_2_label.config(text="File 2: None")
            if hasattr(self,'mfg_pn_var_1'): self.mfg_pn_var_1.set("")
            if hasattr(self,'mfg_pn_var_2'): self.mfg_pn_var_2.set("")
            if hasattr(self,'model_combobox'): self.model_combobox.set(self.placeholder_text); self.model_combobox.state(["disabled"])
            if os.path.exists(self.temp_image_dir):
                try: shutil.rmtree(self.temp_image_dir); print(f"DEBUG: Deleted temp dir: {self.temp_image_dir}")
                except OSError as e: print(f"Error deleting temp dir {self.temp_image_dir}: {e}")
            self._create_temp_image_dir()

        self.pending_user_image_path = None
        self.pending_user_image_pil = None

        if hasattr(self,'conversation_history'): self.conversation_history.config(state=tk.NORMAL); self.conversation_history.delete(1.0, tk.END)
        if hasattr(self,'comparison_treeview'):
            for item in self.comparison_treeview.get_children(): self.comparison_treeview.delete(item)
        self.conversation_log=[]; self.ai_history=[]
        if clear_files: self.update_conversation_history("System: Welcome! Load PDFs to start.",role="system")
        if hasattr(self,'user_input_entry'): self.user_input_entry.delete(0,tk.END)
        self.model=None; self.chat_session=None
        self._configure_ai()
        if not clear_files and self.spec_sheet_1_path and self.spec_sheet_2_path:
            if hasattr(self,'model_combobox'): self.model_combobox.config(state='readonly')
            self.update_conversation_history("System: AI context cleared. Files remain. Select model.",role="system")
        elif clear_files:
             if hasattr(self,'model_combobox'): self.model_combobox.state(['disabled'])
        if hasattr(self, 'start_comparison_button'): self.start_comparison_button.config(state=tk.DISABLED)
        self._update_ui_for_ai_status(api_key_configured=self.api_key_configured,model_initialized=False)
        print("DEBUG: Clear All finished.")

    def extract_text_from_pdf(self, filepath):
        if not filepath or not os.path.exists(filepath): self.update_conversation_history(f"System: PDF not found: {os.path.basename(filepath or 'Unknown')}", role="error"); return ""
        try:
            self.update_conversation_history(f"System: Extracting text from {os.path.basename(filepath)}...", role="system")
            with fitz.open(filepath) as doc: text = "".join(page.get_text() for page in doc)
            self.update_conversation_history(f"System: Text extraction OK: {os.path.basename(filepath)}.", role="system"); return text
        except Exception as e: self.update_conversation_history(f"System: Error extracting text from {os.path.basename(filepath)}: {e}", role="error"); return ""

    def extract_images_from_pdf(self, filepath, output_folder):
        if not filepath or not os.path.exists(filepath): self.update_conversation_history(f"System: PDF not found: {os.path.basename(filepath or 'Unknown')}", role="error"); return []
        paths = []
        try:
            self.update_conversation_history(f"System: Extracting images from {os.path.basename(filepath)}...", role="system")
            if not os.path.exists(output_folder): os.makedirs(output_folder)
            with fitz.open(filepath) as doc:
                for i, page in enumerate(doc):
                    for j, img_info in enumerate(page.get_images(full=True)):
                        xref = img_info[0]
                        try: base = doc.extract_image(xref)
                        except Exception as e: self.update_conversation_history(f"System: Error extracting img xref {xref} pg {i+1}. Skip. Err: {e}", role="error"); continue
                        img_bytes, ext = base["image"], base["ext"]
                        path = os.path.join(output_folder, f"pg{i+1}_img{j+1}.{ext}")
                        try:
                            with open(path, "wb") as f: f.write(img_bytes)
                            paths.append(path)
                        except IOError as e: self.update_conversation_history(f"System: IOError saving image {path}. Error: {e}", role="error")
            msg = f"System: Extracted {len(paths)} images from {os.path.basename(filepath)}." if paths else f"System: No images found in {os.path.basename(filepath)}."
            self.update_conversation_history(msg, role="system"); return paths
        except Exception as e: self.update_conversation_history(f"System: Error extracting images from {os.path.basename(filepath)}: {e}", role="error"); return []

    def check_and_process_spec_sheets(self):
        if not (self.spec_sheet_1_path and self.spec_sheet_2_path): return
        if not self.api_key_configured: self.update_conversation_history("System: API Key not configured.", role="error"); return
        if not self.model: self.update_conversation_history("System: AI Model not selected. Please select a model.", role="system"); return
        self.update_conversation_history("System: Files and model active. Clearing old results...", role="system")
        self.conversation_log = []; self.ai_history = []
        if hasattr(self, 'conversation_history'): self.conversation_history.config(state=tk.NORMAL); self.conversation_history.delete(1.0, tk.END)
        if hasattr(self, 'comparison_treeview'):
            for item in self.comparison_treeview.get_children(): self.comparison_treeview.delete(item)
        if self.api_key_configured: self.update_conversation_history("System: AI Configured.", role="system")
        if self.model: self.update_conversation_history(f"System: Model '{self.model.model_name}' active.", role="system")
        self.process_spec_sheets()

    def process_spec_sheets(self):
        if not self.model or not self.api_key_configured or not self.spec_sheet_1_path or not self.spec_sheet_2_path:
            self.update_conversation_history("System: Pre-reqs not met (files, API key, model).", role="error"); return
        self.update_conversation_history("System: Starting initial analysis...", role="system")
        self.spec_sheet_1_text = self.extract_text_from_pdf(self.spec_sheet_1_path)
        if not self.spec_sheet_1_text: self.update_conversation_history(f"System: Halting. Text extract fail: {os.path.basename(self.spec_sheet_1_path)}.", role="error"); return
        s1_f = os.path.join(self.temp_image_dir, f"{os.path.splitext(os.path.basename(self.spec_sheet_1_path))[0]}_imgs_{len(os.listdir(self.temp_image_dir))}")
        self.spec_sheet_1_image_paths = self.extract_images_from_pdf(self.spec_sheet_1_path, s1_f)
        self.spec_sheet_2_text = self.extract_text_from_pdf(self.spec_sheet_2_path)
        if not self.spec_sheet_2_text: self.update_conversation_history(f"System: Halting. Text extract fail: {os.path.basename(self.spec_sheet_2_path)}.", role="error"); return
        s2_f = os.path.join(self.temp_image_dir, f"{os.path.splitext(os.path.basename(self.spec_sheet_2_path))[0]}_imgs_{len(os.listdir(self.temp_image_dir))}")
        self.spec_sheet_2_image_paths = self.extract_images_from_pdf(self.spec_sheet_2_path, s2_f)

        initial_analysis_prompt_text = (
            "You are an expert electronics component analyst. Analyze the following two component specification sheets.\n\n"
            "**Instructions for AI:**\n"
            "1. For Component 1 (described first), identify its specific component type.\n"
            "2. For Component 2 (described second), identify its specific component type.\n"
            "3. Assess if Component 1 and Component 2 are functionally similar (e.g., both are dual N-channel MOSFETs, or one is an LDO regulator and the other a switching regulator, or one is a TVS diode and the other a Zener diode). Your assessment should be based on their primary function.\n"
            "4. For Component 1, find and extract the first complete Manufacturer Part Number (MFG P/N) listed in its 'Order Information' or equivalent section. If multiple are listed, provide only the first one. If none is explicitly found, state 'Not Found'.\n"
            "5. For Component 2, find and extract the first complete Manufacturer Part Number (MFG P/N) listed in its 'Order Information' or equivalent section. If multiple are listed, provide only the first one. If none is explicitly found, state 'Not Found'.\n\n"
            "**Output Format:**\n"
            "Please provide your response *only* in the following structured format, using these exact labels:\n"
            "Component1_Type: [Type for component 1]\n"
            "Component2_Type: [Type for component 2]\n"
            "Functionally_Similar: [Yes/No, brief explanation]\n"
            "MFG_PN1: [MFG P/N for component 1 or 'Not Found']\n"
            "MFG_PN2: [MFG P/N for component 2 or 'Not Found']\n\n"
            "**Component 1 Data:**\n"
            f"Text Content:\n{self.spec_sheet_1_text}\n\n"
            "**Component 2 Data:**\n"
            f"Text Content:\n{self.spec_sheet_2_text}\n"
        )

        prompt_parts_for_genai = [initial_analysis_prompt_text]
        # Add images for component 1
        for img_path in self.spec_sheet_1_image_paths:
            try: prompt_parts_for_genai.append(Image.open(img_path))
            except Exception as e: self.update_conversation_history(f"System: Error loading image {img_path} for Comp 1. Skip. Err: {e}", role="error")
        # Add images for component 2
        prompt_parts_for_genai.append("\n--- End of Component 1 Images, Start of Component 2 Images (if any) ---") # Separator for clarity if needed
        for img_path in self.spec_sheet_2_image_paths:
            try: prompt_parts_for_genai.append(Image.open(img_path))
            except Exception as e: self.update_conversation_history(f"System: Error loading image {img_path} for Comp 2. Skip. Err: {e}", role="error")

        user_prompt_for_history_log = "User: Initial component type identification and MFG P/N extraction for spec sheets."
        self.send_to_ai(prompt_parts_for_genai, is_initial_analysis=True, user_prompt_for_history=user_prompt_for_history_log)


    def send_to_ai(self, prompt_parts, is_initial_analysis=False, user_prompt_for_history=None):
        if not self.model: self.update_conversation_history("System: AI model N/A.", role="error"); return None
        active_model_name = self.model.model_name
        raw_ai_response_text = ""

        final_prompt_parts = list(prompt_parts) # Work with a copy
        if hasattr(self, 'root'): self.root.update_idletasks()
        if self.translate_to_chinese_var.get():
            translation_instruction = " Please provide your entire response in Chinese."
        else:
            translation_instruction = " Please provide your entire response in English."

        appended_to_text = False
        for i in range(len(final_prompt_parts) - 1, -1, -1):
            if isinstance(final_prompt_parts[i], str):
                final_prompt_parts[i] += translation_instruction
                appended_to_text = True; break
        if not appended_to_text: final_prompt_parts.append(translation_instruction)

        try:
            self.send_button.config(state=tk.DISABLED); self.user_input_entry.config(state=tk.DISABLED)
            if hasattr(self, 'start_comparison_button'): self.start_comparison_button.config(state=tk.DISABLED)
            self.update_conversation_history(f"System: Sending to AI ({active_model_name})... May take time.", role="system")

            if is_initial_analysis and user_prompt_for_history:
                 self._add_to_ai_history('user', user_prompt_for_history)

            response = self.model.generate_content(final_prompt_parts, request_options={'timeout': 600})

            if response.prompt_feedback and response.prompt_feedback.block_reason:
                raw_ai_response_text = f"AI Error - Prompt was blocked. Reason: {response.prompt_feedback.block_reason}"
                self.update_conversation_history(f"System: {raw_ai_response_text}", role="error")
            elif not response.candidates or not hasattr(response, 'text') or not response.text:
                raw_ai_response_text = "AI response empty/no content."
                self.update_conversation_history(f"System: AI ({active_model_name}): {raw_ai_response_text}", role="system")
            else:
                raw_ai_response_text = response.text
                self.update_conversation_history(f"AI ({active_model_name}): {raw_ai_response_text}", role="ai") # Display formatted

            self._add_to_ai_history('model', raw_ai_response_text) # Log model's raw response or error

            if is_initial_analysis:
                self.chat_session = None
                # Parse the response and update UI elements
                parsed_info = self._parse_initial_analysis_response(raw_ai_response_text)

                self.update_conversation_history(f"System: Initial Analysis Parsed Data:", role="system")
                self.update_conversation_history(f"  Component 1 Type: {parsed_info['component1_type']}", role="system")
                self.update_conversation_history(f"  Component 2 Type: {parsed_info['component2_type']}", role="system")
                self.update_conversation_history(f"  Functionally Similar: {parsed_info['functionally_similar']}", role="system")

                if hasattr(self, 'mfg_pn_var_1'):
                    self.mfg_pn_var_1.set(parsed_info['mfg_pn1'] if parsed_info['mfg_pn1'] != "Not Found" else "")
                    self.update_conversation_history(f"  MFG P/N 1 set to: {self.mfg_pn_var_1.get() or 'Not Found'}", role="system")
                if hasattr(self, 'mfg_pn_var_2'):
                    self.mfg_pn_var_2.set(parsed_info['mfg_pn2'] if parsed_info['mfg_pn2'] != "Not Found" else "")
                    self.update_conversation_history(f"  MFG P/N 2 set to: {self.mfg_pn_var_2.get() or 'Not Found'}", role="system")

                if hasattr(self, 'root'): self.root.update_idletasks()

                # Force update Entry widgets UI
                if hasattr(self, 'mfg_pn_entry_1'):
                    self.mfg_pn_entry_1.delete(0, tk.END)
                    self.mfg_pn_entry_1.insert(0, self.mfg_pn_var_1.get())

                if hasattr(self, 'mfg_pn_entry_2'):
                    self.mfg_pn_entry_2.delete(0, tk.END)
                    self.mfg_pn_entry_2.insert(0, self.mfg_pn_var_2.get())

                if hasattr(self, 'start_comparison_button'):
                    if parsed_info["is_similar_flag"] and not ("AI Error" in raw_ai_response_text or "empty/no content" in raw_ai_response_text) :
                        self.start_comparison_button.config(state=tk.NORMAL)
                        self.update_conversation_history("System: Components appear functionally similar. 'Start Detailed Comparison' enabled.", role="system")
                    else:
                        self.start_comparison_button.config(state=tk.DISABLED)
                        self.update_conversation_history("System: Components may not be functionally similar or analysis incomplete. Detailed comparison not enabled.", role="system")
            return raw_ai_response_text
        except Exception as e:
            err_msg = f"System: Error with AI ({active_model_name}): {e}"
            self.update_conversation_history(err_msg, role="error"); print(f"DEBUG: {err_msg}")
            self._add_to_ai_history('model', f"Error: {e}")
            if hasattr(self, 'start_comparison_button'): self.start_comparison_button.config(state=tk.DISABLED)
            if isinstance(e, (google_exceptions.PermissionDenied,google_exceptions.Unauthenticated)): self.api_key_configured=False
            if isinstance(e, (google_exceptions.InvalidArgument, ValueError, BlockedPromptException, StopCandidateException, google_exceptions.NotFound, google_exceptions.PermissionDenied)):
                self.update_conversation_history(f"System: Resetting model ({active_model_name}) due to error.", role="system")
                self.model = None; self.chat_session = None; self.ai_history = []
            return f"AI Error: {e}"
        finally: self._update_ui_for_ai_status()

def main():
    try:
        s = ttk.Style(); available_themes = s.theme_names()
        if "clam" in available_themes: s.theme_use("clam")
        elif "aqua" in available_themes: s.theme_use("aqua")
        elif "vista" in available_themes: s.theme_use("vista")
    except Exception: pass
    root = tk.Tk()
    app = ComponentComparatorAI(root)
    root.mainloop()

if __name__ == "__main__":
    main()

#[end of main.py]
